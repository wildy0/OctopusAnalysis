# Convert xls/x or csv spreadsheet export from octopus smartmeter data to export to analyse it
# Created by Dr Tim Wilding,  2022
# Copyright (c) 2022, Dr Tim Wilding
# All rights reserved.
#
# This source code is licensed under the BSD-style license found in the
# LICENSE file in the root directory of this source tree.
import getopt
import os
import pathlib
from pathlib import Path
from tkinter import filedialog
import tkinter
import sys
import atexit
import pandas as pd
from datetime import datetime
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def table_analysis(idoc, table_data, table_data_rem):
    doctable(idoc, table_data.groupby(['theyear'])['energy_use'].sum().reset_index(), 'Yearly Use')

    day_data = table_data_rem.groupby(['yearday', 'theyear', 'thewday', 'theday', 'themonth'])[
        'energy_use'].sum().reset_index()

    # daily
    themean = day_data.groupby(['theyear'])['energy_use'].mean().reset_index(name='Mean').set_index('theyear')
    themax = day_data.groupby(['theyear'])['energy_use'].max().reset_index(name='Max').set_index('theyear')
    themin = day_data.groupby(['theyear'])['energy_use'].min().reset_index(name='Min').set_index('theyear')
    d_use = pd.concat([themean, themax, themin], axis=1).reset_index()
    doctable(idoc, d_use, 'Daily Use')
    # hourly
    hrn_data = table_data_rem.groupby(['theyear', 'yearday', 'thewday', 'thehour'])['energy_use'].sum().reset_index()
    themean = hrn_data.groupby(['theyear'])['energy_use'].mean().reset_index(name='Mean').set_index('theyear')
    themax = hrn_data.groupby(['theyear'])['energy_use'].max().reset_index(name='Max').set_index('theyear')
    themin = hrn_data.groupby(['theyear'])['energy_use'].min().reset_index(name='Min').set_index('theyear')
    d_use = pd.concat([themean, themax, themin], axis=1).reset_index()
    doctable(idoc, d_use, 'Hourly Use')
    doctable(idoc, hrn_data[hrn_data["thehour"].between(0, 6, inclusive='both')].groupby(['theyear'])
                ['energy_use'].sum().reset_index(), 'Night Use 00:00 - 06:00')
    doctable(idoc, hrn_data[hrn_data["thehour"].between(16, 18, inclusive='both')].groupby(['theyear'])
                ['energy_use'].sum().reset_index(), 'Total use 16:00 to 18:00')


def doctable(idoc, data, tabletitle):
    idoc.add_heading(tabletitle)
    table = idoc.add_table(rows=(data.shape[0]+1), cols=data.shape[1])
    table.allow_autofit = True
    table.autofit = True
    i = 0
    for name in data.keys():
        table.cell(0, i).text = name
        i += 1
    for i, column in enumerate(data):
        for row in range(data.shape[0]):
            table.cell(row+1, i).text = "{:.1f}".format((data.iloc[row][column]))


def do_hour_plot(data):
    hrn_data = data.groupby(['theyear', 'yearday', 'thewday', 'thehour'])['energy_use'].sum().reset_index()
    hr_data = hrn_data.groupby(['theyear', 'thehour'])['energy_use'].mean().reset_index()
    hr_data_min = hrn_data.groupby(['theyear', 'thehour'])['energy_use'].min().reset_index()
    hr_data_max = hrn_data.groupby(['theyear', 'thehour'])['energy_use'].max().reset_index()

    for year in hr_data['theyear'].unique():
        u = hr_data[hr_data["theyear"] == year]['energy_use'].to_numpy()
        u_min = u - hr_data_min[hr_data["theyear"] == year]['energy_use'].to_numpy()
        u_max = hr_data_max[hr_data["theyear"] == year]['energy_use'].to_numpy() - u
        yerr = [u_min, u_max]
        plt.errorbar(hr_data[hr_data["theyear"] == year]['thehour'],
                     u, yerr=yerr, marker='o', solid_capstyle='projecting', capsize=5,
                     label=str(year))


def do_the_month_plot(data, plt_label, show_plot=True, save_plot=False, output_path="", file_prefix=""):
    fig = plt.figure(figsize=(16, 12), dpi=150)
    plt.title('Energy use (months)')
    plt.suptitle(plt_label)
    pnum = 1
    for mth in range(1, 13):
        # do plot
        ax = plt.subplot(7, 2, pnum)
        ax.title.set_text('Average by hour Month {}'.format(mth))
        do_hour_plot(data[data['themonth'] == mth])
        handles, labels = ax.get_legend_handles_labels()
        # reverse the order
        ax.legend(handles[::-1], labels[::-1], loc="upper right")
        pnum += 1
    # now do winter and summer
    # winter
    ax = plt.subplot(7, 2, pnum)
    ax.title.set_text('Average by hour Winter')
    do_hour_plot(data[(((data['themonth'] >= 1) & (data['themonth'] <= 2)) | (data['themonth'] == 12))])
    handles, labels = ax.get_legend_handles_labels()
    # reverse the order
    ax.legend(handles[::-1], labels[::-1], loc="upper right")
    pnum += 1
    # summer
    ax = plt.subplot(7, 2, pnum)
    ax.title.set_text('Average by hour Summer')
    do_hour_plot(data[data['themonth'].between(6, 8, inclusive='both')])
    handles, labels = ax.get_legend_handles_labels()
    # reverse the order
    ax.legend(handles[::-1], labels[::-1], loc="upper right")
    fig.subplots_adjust(hspace=0.5)
    plot_save_name = os.path.join(output_path, file_prefix + "_" + plt_label + ".png")
    if save_plot:
        plt.savefig(plot_save_name, dpi=200)
    if not show_plot:
        fig.clear()
        plt.close(fig)
    return plot_save_name


def do_the_plot(data, plt_label, show_plot=True, save_plot=False, output_path="", file_prefix=""):
    d_data = data.groupby(['theyear', 'yearday'])['energy_use'].sum().reset_index()
    d_data_m = data.groupby(['theyear', 'themonth', 'yearday'])['energy_use'].sum().reset_index()
    ad_data = d_data_m.groupby(['theyear', 'themonth'])['energy_use'].mean().reset_index()
    ad_data_min = d_data_m.groupby(['theyear', 'themonth'])['energy_use'].min().reset_index()
    ad_data_max = d_data_m.groupby(['theyear', 'themonth'])['energy_use'].max().reset_index()
    m_data = data.groupby(['theyear', 'themonth'])['energy_use'].sum().reset_index()

    fig = plt.figure(figsize=(16, 12), dpi=150)

    plt.title('Energy use')
    plt.suptitle(plt_label)
    ax = plt.subplot(4, 1, 1)
    ax.title.set_text('Daily use')
    for year in d_data['theyear'].unique():
        plt.plot("yearday", "energy_use", data=d_data[d_data["theyear"] == year], linestyle='-',
                 marker='o', label=str(year))

    handles, labels = ax.get_legend_handles_labels()
    # reverse the order
    ax.legend(handles[::-1], labels[::-1], loc="upper right")

    ax2 = plt.subplot(4, 1, 2)
    ax2.title.set_text('Monthly use')
    for year in m_data['theyear'].unique():
        plt.plot("themonth", "energy_use", data=m_data[m_data["theyear"] == year], linestyle='-',
                 marker='o', label=str(year))

    handles, labels = ax2.get_legend_handles_labels()
    # reverse the order
    ax2.legend(handles[::-1], labels[::-1], loc="upper right")

    ax3 = plt.subplot(4, 1, 3)
    ax3.title.set_text('Average daily use by month')
    for year in ad_data['theyear'].unique():
         u = ad_data[ad_data["theyear"] == year]['energy_use'].to_numpy()
         u_min = u - ad_data_min[ad_data["theyear"] == year]['energy_use'].to_numpy()
         u_max = ad_data_max[ad_data["theyear"] == year]['energy_use'].to_numpy() - u
         yerr = [u_min, u_max]
         plt.errorbar(ad_data[ad_data["theyear"] == year]['themonth'],
                      u, yerr=yerr, marker='o', solid_capstyle='projecting', capsize=5, label=str(year))
    handles, labels = ax3.get_legend_handles_labels()
    # reverse the order
    ax3.legend(handles[::-1], labels[::-1], loc="upper right")

    ax4 = plt.subplot(4, 1, 4)
    ax4.title.set_text('Average by hour')
    do_hour_plot(big_data)
    handles, labels = ax4.get_legend_handles_labels()
    # reverse the order
    ax4.legend(handles[::-1], labels[::-1], loc="upper right")
    fig.subplots_adjust(hspace=0.4)
    plot_save_name = os.path.join(output_path, file_prefix + "_" + plt_label + ".png")

    if save_plot:
        plt.savefig(plot_save_name, dpi=200)

    if not show_plot:
        fig.clear()
        plt.close(fig)
    return plot_save_name


if __name__ == '__main__':
    atexit.register(input, "Enter any Key to Close/Exit")

    global nodelete
    nodelete = False

    start_key = ' Start'
    end_key = ' End'
    power_key_prefix = 'Consumption'
    power_key = ''
    electric_unit = "kWh"
    gas_calorific = 1.02264 * 40.0 / 3.6
    time_format = '%Y-%m-%dT%H:%M:%S%z'


    try:
        opts, args = getopt.getopt(sys.argv[1:], "hn", ["help", "nodelete"])
        for opt, arg in opts:
            print(opt, arg)
            if opt in ('-h', "--help"):
                print("Valid options: -h or --help show help\n-n or --nodelete do not remove missing data.")
                sys.exit()
            elif opt in ("-n", "--nodelete"):
                nodelete = True
                print("No delete option set.")
    except getopt.GetoptError:
        print("Valid options: -h or --help show help\n-n or --nodelete do not remove missing data.")
        sys.exit()

    print("Energy analysis.")

    # start a root tk window and hide it now, so that it goes away before we do the filedialog
    root = tkinter.Tk()
    root.withdraw()
    # root.update()

    filename = filedialog.askopenfilename(title="Select Syllabus plus Excel spreadsheet export by year,"
                                                " or spreadsheet output for drafts formatted as generated by "
                                                "this script",
                                          # filetypes=(("Excel sheet", "*.xls*"), ("all", "*.*")))
                                          filetypes=[("Excel sheet", ".xlsx .xls .csv")]
                                          )

    print("Reading spreadsheet file %s" % filename)

    fullpath = Path(filename)
    filepath = str(fullpath.parent)
    # remove spaces and . from filename when creating output path to avoid issues with directories with spaces or
    # multiple .
    f_type = fullpath.suffix
    i_filename = "_".join(fullpath.stem.split(" "))
    i_filename = "_".join(i_filename.split("."))
    # make the path
    out_dir = pathlib.PurePath(filepath, i_filename)

    print("directory: %s name: %s" % (filepath, i_filename))

    if f_type == '.csv':
        try:
            df = pd.read_csv(filename)  # Read a csv file
        except PermissionError:
            print("Could not open the file, try close it if open and check permissions before trying again.\n")
            sys.exit(1)
        except AssertionError:
            print("Please select a file.  Exiting.\n")
            sys.exit(1)
    else:
        try:
            file = pd.ExcelFile(filename)  # Establishes the Excel file you wish to import into Pandas
        except PermissionError:
            print("Could not open the file, try close it if open and check permissions before trying again.\n")
            sys.exit(1)
        except AssertionError:
            print("Please select a file.  Exiting.\n")
            sys.exit(1)
        sheet_map = pd.read_excel(file, sheet_name=None)
        df = sheet_map[list(sheet_map.keys())[0]]

    for key in df.keys():
        if power_key_prefix in key:
            power_key = key
    # first_column_key = df.keys()[0]

    if power_key == '':
        print("Could not find the data column {}".format(power_key_prefix))
        print("Wrong file type, or wrong file type options.")
        sys.exit(1)

    if start_key not in df.keys():
        print("Could not find the data column {}".format(start_key))
        print("Wrong file type, or wrong file type options.")
        sys.exit(1)

    if end_key not in df.keys():
        print("Could not find the data column {}".format(end_key))
        print("Wrong file type, or wrong file type options.")
        sys.exit(1)

    number_of_rows = len(df.index)
    print("Processing {} data rows in the file.  This will take a bit of time, "
          "but not long enough to grab a coffee, or do anything much really. "
          " It is pretty fast.".format(number_of_rows))

    if power_key.find(electric_unit) == -1:
        gas_file = 1
        print("Reading a Gas File")
        print("Converting to KWh using calorific value: {}".format(gas_calorific))
        df['energy_use'] = df[power_key] * gas_calorific
    else:
        gas_file = 0
        print("Reading a Electricity File")
        df['energy_use'] = df[power_key]

    # there is probably a more optimal way to do this but life is too short for optimisation
    print("Please wait")
    df['thetime'] = df.apply(lambda row: datetime.strptime(row[start_key].lstrip(), time_format), axis=1)
    df['thestoptime'] = df.apply(lambda row: datetime.strptime(row[end_key].lstrip(), time_format), axis=1)
    df['duration'] = df.apply(lambda row: (row['thestoptime'] - row['thetime']).seconds//60, axis=1)
    print("Please wait")
    df['thewday'] = df.apply(lambda row: datetime.strptime(row[start_key].lstrip(), time_format).weekday(),
                             axis=1)
    print("Please wait")
    df['theday'] = df.apply(lambda row: datetime.strptime(row[start_key].lstrip(), time_format).day, axis=1)
    print("Please wait")
    df['themonth'] = df.apply(lambda row: datetime.strptime(row[start_key].lstrip(), time_format).month,
                              axis=1)
    print("Please wait a bit longer")
    df['theyear'] = df.apply(lambda row: datetime.strptime(row[start_key].lstrip(), time_format).year, axis=1)
    print("Please wait almost done, half way now")
    df['thehour'] = df.apply(lambda row: datetime.strptime(row[start_key].lstrip(), time_format).hour, axis=1)
    print("Please wait a bit longer")
    df['theweek'] = df.apply(lambda row: datetime.strptime(row[start_key].lstrip(), time_format).
                             isocalendar()[1], axis=1)
    print("Please wait a bit longer")
    df['yearday'] = df.apply(lambda row: datetime.strptime(row[start_key].lstrip(), time_format).timetuple().
                             tm_yday, axis=1)
    df['uday'] = df.apply(lambda row: ((row['theyear'] * 1000) + row['yearday']), axis=1)
    print('Finished processing the data')
    big_data = df

    print("\nNow checking data and creating plots and report.")
    pdf_string = ""
    pdf_string += "The script looks for days with missing data during the day. " \
                  "I have found around 1 to 20% data have missing hour data." \
                  "I assume smart meter data is lost from time to time,  particularly if you have power cuts?\n"

    durations = (big_data['duration'].groupby(big_data['uday']).sum())  # calculate the duration of every row time point
    durations_t = (big_data['duration'].groupby(big_data['uday']).transform('sum'))
    total_data_days = len(durations)
    print("{} days".format(total_data_days))
    missing_data_days = np.sum(durations.to_numpy() < 24*60)  # look for days with less than 24-hours of data
    percentage_missing = missing_data_days/total_data_days*100

    if missing_data_days > 0:
        print("The energy file contains missing data.")
        print("This will be removed from calculations of min/max/mean daily/hourly use.")
        if not nodelete:
            print("Removed {} days of {} days ({:.0f}%).\n".format(missing_data_days, total_data_days,
                                                                                    percentage_missing))
            pdf_string += "Removed {} days of {} days ({:.0f}%) because of missing data.\n".format(missing_data_days,
                                                                                    total_data_days, percentage_missing)
            big_data_removed = big_data[durations_t >= 24*60]  # only keep days with 24-hours of data in them
        else:
            pdf_string += "{} days of {} data days ({:.0f}%) had missing data but " \
                          "were NOT REMOVED as nodelete option used.\n".\
                format(missing_data_days, total_data_days, percentage_missing)
            print("{} data points of {} data points ({:.0f}%) had missing data "
                  "but were NOT REMOVED as nodelete option used\n".
                  format(missing_data_days, total_data_days, percentage_missing))
            big_data_removed = big_data
    else:
        print("The energy file did not contain missing data.")
        pdf_string += "{} days".format(total_data_days)
        pdf_string += "No missing data was found all days were complete\n"
        big_data_removed = big_data

    print("Creating Report")
    wdocument = Document()
    for line in pdf_string.split('\n'):
        p = wdocument.add_paragraph()
        r = p.add_run()
        r.add_text(line)

    wdocument.add_heading('All data')
    table_analysis(wdocument, big_data, big_data_removed)

    wdocument.add_heading('Winter')

    table_analysis(wdocument, big_data[((big_data['themonth'] >= 1) & (big_data['themonth'] <= 2))
                                | (big_data['themonth'] == 12)], big_data_removed[((big_data_removed['themonth'] >= 1)
                                 & (big_data_removed['themonth'] <= 2)) | (big_data_removed['themonth'] == 12)])

    wdocument.add_heading('Summer')
    table_analysis(wdocument, big_data[big_data['themonth'].between(6, 8, inclusive='both')], big_data_removed
                                            [big_data_removed['themonth'].between(6, 8, inclusive='both')])

    plot_names = []
    plot_names.append(do_the_plot(big_data_removed, "All", save_plot=True,
                                  output_path=filepath, file_prefix=i_filename))
    plot_names.append(do_the_month_plot(big_data_removed, "Months", save_plot=True,
                                        output_path=filepath, file_prefix=i_filename))

    for p_name in plot_names:
        print("Adding plot image {} to report.".format(p_name))
        ip = wdocument.add_paragraph()
        pr = ip.add_run()
        pr.add_picture(p_name, width=Inches(6))
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rep_name = os.path.join(filepath, i_filename + ".docx")
    wdocument.save(rep_name)
    print("Report saved to {}".format(rep_name))
